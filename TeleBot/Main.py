import telebot
import docx
from docx import Document

bot = telebot.TeleBot('5608205852:AAFnDjh4XtJiwx0I-20vgo_WReaLl5dD0EA')
#xxx = Document(docx=None)
#doc = Document()
#doc.add_heading("", 0)
#doc.save("doc.docx")




@bot.message_handler(commands=["start"])
def start(message):
    name, lname = message.from_user.first_name, message.from_user.last_name
    bot.send_message(message.chat.id, f"Hi, {name} {lname}")

@bot.message_handler(commands=["qqq"])
def qqq(message):
    file = open('wise_oak.docx', 'rb')
    bot.send_document(message.chat.id, file)


@bot.message_handler(commands=['resume'])
def resume(message):
    bot.send_message(message.chat.id, "Enter your name")
    mode = 'resume'


@bot.message_handler()
def get_text(message):
    text = message.text


bot.polling(non_stop=True)